import os
import time
import uuid
import random
import shutil
import logging 
import subprocess
from pathlib import Path
from flask import (
    Flask, request, render_template, send_file, flash, redirect, url_for
)
from werkzeug.utils import secure_filename

from PyPDF2 import PdfMerger, PdfReader, PdfWriter
from PIL import Image, ImageFile
ImageFile.LOAD_TRUNCATED_IMAGES = True
import pdfplumber
from pdf2image import convert_from_path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import zipfile
import fitz  # PyMuPDF

# --- CONFIGURATION ---
BASE_DIR = Path(__file__).parent.resolve()
UPLOAD_FOLDER = BASE_DIR / "uploads"
TEMP_FOLDER = BASE_DIR / "temp"
for folder in [UPLOAD_FOLDER, TEMP_FOLDER]:
    folder.mkdir(parents=True, exist_ok=True)

ALLOWED_EXTENSIONS = {
    'pdf': ['pdf'],
    'word': ['docx', 'doc'],
    'image': ['jpg', 'jpeg', 'png']
}

# --- FLASK APP ---
app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get("SECRET_KEY") or os.urandom(24)
app.config['UPLOAD_FOLDER'] = str(UPLOAD_FOLDER)
app.config['TEMP_FOLDER'] = str(TEMP_FOLDER)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB

# --- LOGGING ---
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s %(levelname)s %(message)s',
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# --- UTILITIES ---
def allowed_file(filename, file_type):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS.get(file_type, [])

def get_unique_filename(filename):
    name, ext = os.path.splitext(filename)
    return f"{name}_{uuid.uuid4().hex[:8]}{ext}"

def libreoffice_available():
    """Check if LibreOffice or soffice is available in PATH."""
    return shutil.which("libreoffice") or shutil.which("soffice")

def libreoffice_convert_to_pdf(input_path, output_path):
    """
    Converts a Word file to PDF using LibreOffice in headless mode.
    """
    try:
        libreoffice_cmd = shutil.which("libreoffice") or shutil.which("soffice")
        if not libreoffice_cmd:
            raise RuntimeError("LibreOffice/soffice not found in PATH.")
        subprocess.check_call([
            libreoffice_cmd,
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', str(output_path.parent),
            str(input_path)
        ])
        # LibreOffice names output as <input_basename>.pdf
        output_pdf = output_path.parent / (input_path.stem + '.pdf')
        if not output_pdf.exists():
            raise RuntimeError("LibreOffice failed to create output PDF.")
        return output_pdf
    except Exception as e:
        raise RuntimeError(f"LibreOffice conversion failed: {e}")

def plaintext_docx_to_pdf(docx_path, pdf_path):
    """
    Fallback: Converts a DOCX to a simple text-based PDF using python-docx and reportlab.
    Only text will be preserved.
    """
    doc = Document(docx_path)
    pdf = SimpleDocTemplate(pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    for para in doc.paragraphs:
        if para.text.strip():
            story.append(Paragraph(para.text, styles['Normal']))
            story.append(Spacer(1, 12))
    pdf.build(story)
    return pdf_path

def clean_temp_files():
    now = time.time()
    for folder in [UPLOAD_FOLDER, TEMP_FOLDER]:
        for item in Path(folder).iterdir():
            try:
                if item.is_file() and (now - item.stat().st_mtime) > 3600:
                    item.unlink()
            except Exception as e:
                logger.error(f"Error cleaning temp file {item}: {e}")

@app.before_request
def periodic_cleanup():
    if random.random() < 0.025:  # Run cleanup ~2.5% of requests
        clean_temp_files()

@app.errorhandler(413)
def file_too_large(e):
    flash("File too large! Maximum file size is 16MB.", "error")
    return redirect(url_for('index'))

@app.errorhandler(500)
def internal_error(e):
    flash("An unexpected error occurred. Please try again.", "error")
    return redirect(url_for('index'))

# --- ROUTES ---
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/word_to_pdf', methods=['POST'])
def word_to_pdf():
    file = request.files.get('file')
    if not file or file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('index'))
    if not allowed_file(file.filename, 'word'):
        flash('Please upload a Word file.', 'error')
        return redirect(url_for('index'))
    try:
        filename = get_unique_filename(secure_filename(file.filename))
        temp_word = UPLOAD_FOLDER / filename
        file.save(temp_word)
        output_pdf = TEMP_FOLDER / f"{temp_word.stem}.pdf"
        pdf_path = None
        used_fallback = False
        if libreoffice_available():
            try:
                pdf_path = libreoffice_convert_to_pdf(temp_word, output_pdf)
            except Exception as e:
                logger.warning(f"LibreOffice conversion failed, falling back: {e}")
                used_fallback = True
                pdf_path = plaintext_docx_to_pdf(temp_word, output_pdf)
        else:
            used_fallback = True
            pdf_path = plaintext_docx_to_pdf(temp_word, output_pdf)
        temp_word.unlink()
        if used_fallback:
            flash("Notice: Only text was preserved, formatting/images may be lost (LibreOffice not available).", "warning")
        return send_file(pdf_path, as_attachment=True, download_name=pdf_path.name)
    except Exception as e:
        logger.exception(e)
        flash("Conversion failed. Please ensure LibreOffice is installed for best results.", "error")
        return redirect(url_for('index'))

@app.route('/pdf_to_word', methods=['POST'])
def pdf_to_word():
    file = request.files.get('file')
    if not file or file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('index'))
    if not allowed_file(file.filename, 'pdf'):
        flash('Please upload a PDF file.', 'error')
        return redirect(url_for('index'))
    try:
        filename = get_unique_filename(secure_filename(file.filename))
        temp_pdf = UPLOAD_FOLDER / filename
        file.save(temp_pdf)
        doc = Document()
        with pdfplumber.open(temp_pdf) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    doc.add_paragraph(text)
        docx_path = TEMP_FOLDER / (temp_pdf.stem + ".docx")
        doc.save(docx_path)
        temp_pdf.unlink()
        return send_file(docx_path, as_attachment=True, download_name=docx_path.name)
    except Exception as e:
        logger.exception(e)
        flash("Conversion failed. Only basic text will be extracted.", "error")
        return redirect(url_for('index'))

@app.route('/pdf_to_image', methods=['POST'])
def pdf_to_image():
    file = request.files.get('file')
    if not file or file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('index'))
    if not allowed_file(file.filename, 'pdf'):
        flash('Please upload a PDF file.', 'error')
        return redirect(url_for('index'))
    try:
        filename = get_unique_filename(secure_filename(file.filename))
        temp_pdf = UPLOAD_FOLDER / filename
        file.save(temp_pdf)
        images = convert_from_path(str(temp_pdf), dpi=200, fmt='jpeg')
        temp_dir = TEMP_FOLDER / f"pdf2img_{uuid.uuid4().hex[:8]}"
        temp_dir.mkdir(parents=True, exist_ok=True)
        zip_path = TEMP_FOLDER / f"{temp_pdf.stem}_images.zip"
        with zipfile.ZipFile(zip_path, "w") as zipf:
            for i, img in enumerate(images, 1):
                img_path = temp_dir / f"page_{i}.jpg"
                img.save(img_path, "JPEG", quality=90)
                zipf.write(img_path, img_path.name)
        shutil.rmtree(temp_dir)
        temp_pdf.unlink()
        return send_file(zip_path, as_attachment=True, download_name=zip_path.name)
    except Exception as e:
        logger.exception(e)
        flash("PDF to Image conversion failed. Make sure PDF is not encrypted/corrupted.", "error")
        return redirect(url_for('index'))

@app.route('/image_to_pdf', methods=['POST'])
def image_to_pdf():
    file = request.files.get('file')
    if not file or file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('index'))
    if not allowed_file(file.filename, 'image'):
        flash('Please upload a JPG/JPEG/PNG file.', 'error')
        return redirect(url_for('index'))
    try:
        img = Image.open(file.stream)
        if img.mode in ('RGBA', 'LA'):
            background = Image.new('RGB', img.size, (255, 255, 255))
            background.paste(img, mask=img.split()[-1])
            img = background
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        pdf_path = TEMP_FOLDER / f"{Path(file.filename).stem}_{uuid.uuid4().hex[:8]}.pdf"
        img.save(pdf_path, "PDF", resolution=100.0)
        return send_file(pdf_path, as_attachment=True, download_name=pdf_path.name)
    except Exception as e:
        logger.exception(e)
        flash("Image to PDF conversion failed.", "error")
        return redirect(url_for('index'))

@app.route('/compress_pdf', methods=['POST'])
def compress_pdf():
    file = request.files.get('file')
    if not file or file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('index'))
    if not allowed_file(file.filename, 'pdf'):
        flash('Please upload a PDF file.', 'error')
        return redirect(url_for('index'))
    try:
        temp_pdf = UPLOAD_FOLDER / get_unique_filename(secure_filename(file.filename))
        file.save(temp_pdf)
        out_pdf = TEMP_FOLDER / f"{temp_pdf.stem}_compressed.pdf"
        doc = fitz.open(temp_pdf)
        doc.save(out_pdf, deflate=True, garbage=4, clean=True, linear=True)
        doc.close()
        temp_pdf.unlink()
        return send_file(out_pdf, as_attachment=True, download_name=out_pdf.name)
    except Exception as e:
        logger.exception(e)
        flash("Compression failed. Only standard PDFs are supported.", "error")
        return redirect(url_for('index'))

@app.route('/merge_pdfs', methods=['POST'])
def merge_pdfs():
    files = request.files.getlist('files[]')
    if not files or files[0].filename == '':
        flash('No files selected', 'error')
        return redirect(url_for('index'))
    try:
        merger = PdfMerger()
        temp_files = []
        for file in files:
            if not allowed_file(file.filename, 'pdf'):
                continue
            temp_pdf = UPLOAD_FOLDER / get_unique_filename(secure_filename(file.filename))
            file.save(temp_pdf)
            merger.append(str(temp_pdf))
            temp_files.append(temp_pdf)
        out_pdf = TEMP_FOLDER / f"merged_{uuid.uuid4().hex[:8]}.pdf"
        merger.write(str(out_pdf))
        merger.close()
        for tf in temp_files:
            tf.unlink()
        return send_file(out_pdf, as_attachment=True, download_name=out_pdf.name)
    except Exception as e:
        logger.exception(e)
        flash("PDF merging failed. Make sure all files are valid PDFs.", "error")
        return redirect(url_for('index'))

@app.route('/split_pdf', methods=['POST'])
def split_pdf():
    file = request.files.get('file')
    if not file or file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('index'))
    if not allowed_file(file.filename, 'pdf'):
        flash('Please upload a PDF file.', 'error')
        return redirect(url_for('index'))
    try:
        temp_pdf = UPLOAD_FOLDER / get_unique_filename(secure_filename(file.filename))
        file.save(temp_pdf)
        reader = PdfReader(str(temp_pdf))
        temp_dir = TEMP_FOLDER / f"split_{uuid.uuid4().hex[:8]}"
        temp_dir.mkdir(parents=True, exist_ok=True)
        for i, page in enumerate(reader.pages):
            writer = PdfWriter()
            writer.add_page(page)
            split_path = temp_dir / f"page_{i+1}.pdf"
            with open(split_path, "wb") as f:
                writer.write(f)
        zip_path = TEMP_FOLDER / f"{temp_pdf.stem}_split.zip"
        with zipfile.ZipFile(zip_path, "w") as zipf:
            for pdf_file in temp_dir.glob("*.pdf"):
                zipf.write(pdf_file, pdf_file.name)
        shutil.rmtree(temp_dir)
        temp_pdf.unlink()
        return send_file(zip_path, as_attachment=True, download_name=zip_path.name)
    except Exception as e:
        logger.exception(e)
        flash("PDF split failed. Only standard PDFs are supported.", "error")
        return redirect(url_for('index'))

@app.route('/encrypt_pdf', methods=['POST'])
def encrypt_pdf():
    file = request.files.get('file')
    password = request.form.get('password')
    if not file or file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('index'))
    if not password:
        flash('Password is required.', 'error')
        return redirect(url_for('index'))
    if not allowed_file(file.filename, 'pdf'):
        flash('Please upload a PDF file.', 'error')
        return redirect(url_for('index'))
    try:
        temp_pdf = UPLOAD_FOLDER / get_unique_filename(secure_filename(file.filename))
        file.save(temp_pdf)
        reader = PdfReader(str(temp_pdf))
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        writer.encrypt(password)
        out_pdf = TEMP_FOLDER / f"{temp_pdf.stem}_encrypted.pdf"
        with open(out_pdf, "wb") as f:
            writer.write(f)
        temp_pdf.unlink()
        return send_file(out_pdf, as_attachment=True, download_name=out_pdf.name)
    except Exception as e:
        logger.exception(e)
        flash("PDF encryption failed. Only standard PDFs are supported.", "error")
        return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=False, port=5000)